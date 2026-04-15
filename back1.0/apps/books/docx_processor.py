import os
import re
from typing import List, Dict, Any
from docx import Document
import tempfile
import uuid

class DocxProcessor:
    """基于python-docx的.docx文件处理器，无需安装Microsoft Word"""
    
    def __init__(self, file_path: str = None):
        """初始化处理器
        
        Args:
            file_path: .docx文件路径（可选）
        """
        self.file_path = file_path
        self.document = None
    
    def open_document(self, file_path: str = None) -> bool:
        """打开.docx文件
        
        Args:
            file_path: .docx文件路径（可选，默认使用初始化时的路径）
            
        Returns:
            成功返回True，失败返回False
        """
        try:
            if file_path:
                self.file_path = file_path
            
            if not self.file_path:
                print("未提供文件路径")
                return False
            
            # 使用python-docx打开文档
            self.document = Document(self.file_path)
            return True
        except Exception as e:
            print(f"打开文档失败: {str(e)}")
            self.close_document()
            return False
    
    def close_document(self):
        """关闭文档（python-docx自动管理资源，这里主要是重置状态）"""
        self.document = None
    
    def extract_toc(self) -> List[Dict[str, Any]]:
        """提取文档目录
        
        Returns:
            目录项列表，每个目录项包含标题、级别和页码
        """
        if not self.document:
            return []
        
        toc_items = []
        
        try:
            # python-docx不直接支持提取目录，我们通过识别标题样式来模拟目录
            # 这里我们使用实际的标题作为目录项
            headings = self._extract_headings()
            
            for i, heading in enumerate(headings):
                toc_item = {
                    'title': heading['text'],
                    'level': heading['level'],
                    'page_num': i + 1  # python-docx无法获取页码，这里用序号代替
                }
                toc_items.append(toc_item)
                
        except Exception as e:
            print(f"提取目录失败: {str(e)}")
        
        return toc_items
    
    def _extract_headings(self) -> List[Dict[str, Any]]:
        """提取文档中的所有标题
        
        Returns:
            标题列表，每个标题包含文本和级别
        """
        headings = []
        
        # 1. 首先尝试通过样式识别标题
        heading_style_map = {
            'Heading 1': 1,
            'Heading 2': 2,
            'Heading 3': 3,
            'Heading 4': 4,
            'Heading 5': 5,
            'Heading 6': 6,
            'heading 1': 1,
            'heading 2': 2,
            'heading 3': 3,
            'heading 4': 4,
            'heading 5': 5,
            'heading 6': 6,
        }
        
        for paragraph in self.document.paragraphs:
            style_name = paragraph.style.name
            if style_name in heading_style_map:
                headings.append({
                    'text': paragraph.text.strip(),
                    'level': heading_style_map[style_name]
                })
        
        # 如果通过样式没有找到标题，尝试通过内容结构识别
        if not headings:
            headings = self._extract_headings_by_content()
        
        return headings
    
    def _extract_headings_by_content(self) -> List[Dict[str, Any]]:
        """通过内容结构识别标题
        
        Returns:
            标题列表，每个标题包含文本和级别
        """
        headings = []
        
        # 定义需要跳过的内容模式
        skip_patterns = [
            re.compile(r'^\d{4}年\d{1,2}月\d{1,2}日$'),  # 日期格式
            re.compile(r'^目\s*录$'),  # 目录标题
            re.compile(r'^\s*$'),  # 空行
        ]
        
        # 章节标题正则表达式模式
        # 匹配类似 "第4章 程序的控制结构" 和 "4.1 程序设计与算法" 这样的标题格式
        chapter_patterns = [
            (re.compile(r'^第\d+章\s+[\w\u4e00-\u9fa5]+'), 1),  # "第4章 程序的控制结构" 格式
            (re.compile(r'^\d+\.\d+\s+[\w\u4e00-\u9fa5]+'), 2),  # "4.1 程序设计与算法" 格式
            (re.compile(r'^\d+\.\d+\.\d+\s+[\w\u4e00-\u9fa5]+'), 3),  # "4.1.1 程序设计的概念" 格式
        ]
        
        for paragraph in self.document.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue
            
            # 跳过不需要的内容
            skip = False
            for pattern in skip_patterns:
                if pattern.match(text):
                    skip = True
                    break
            if skip:
                continue
            
            # 尝试匹配各种标题模式
            for pattern, level in chapter_patterns:
                if pattern.match(text):
                    # 清理标题，去除多余空格和可能的页码
                    cleaned_text = re.sub(r'\s{2,}', ' ', text)
                    cleaned_text = re.sub(r'\s*\d+$', '', cleaned_text)
                    headings.append({
                        'text': cleaned_text,
                        'level': level
                    })
                    break
        
        return headings
    
    def extract_content_by_headings(self) -> List[Dict[str, Any]]:
        """按章节标题提取正文内容
        
        Returns:
            章节内容列表，每个章节包含标题、级别、内容等信息
        """
        if not self.document:
            return []
        
        chapters = []
        
        try:
            # 获取所有段落和表格
            paragraphs = self.document.paragraphs
            tables = self.document.tables
            total_paragraphs = len(paragraphs)
            total_tables = len(tables)
            
            # 直接从段落中识别标题和内容
            current_chapter = None
            
            # 跟踪当前处理的表格索引
            current_table_index = 0
            
            # 跟踪当前处理的段落索引
            current_paragraph_index = 0
            
            # 遍历文档中的所有元素
            while current_paragraph_index < total_paragraphs:
                paragraph = paragraphs[current_paragraph_index]
                text = paragraph.text
                stripped_text = text.strip()
                
                # 跳过目录内容（样式以'toc'开头的段落）
                if paragraph.style.name.startswith('toc') or paragraph.style.name.startswith('TOC'):
                    current_paragraph_index += 1
                    continue
                    
                # 检查是否是标题
                is_heading = False
                heading_level = 0
                
                # 1. 检查样式是否为标题
                style_name = paragraph.style.name
                if style_name.startswith('Heading') or style_name.startswith('heading'):
                    is_heading = True
                    # 从样式名中提取级别
                    try:
                        heading_level = int(style_name[7:])
                    except (ValueError, IndexError):
                        heading_level = 1
                else:
                    # 2. 通过内容结构识别标题
                    # 首先检测主要章节格式（第X章）
                    if re.match(r'^第\d+章\s+[\w\u4e00-\u9fa5]+', stripped_text):
                        is_heading = True
                        heading_level = 1
                    # 检测二级标题格式（X.X）
                    elif re.match(r'^\d+\.\d+\s+[\w\u4e00-\u9fa5]+', stripped_text):
                        is_heading = True
                        heading_level = 2
                    # 检测三级标题格式（X.X.X）
                    elif re.match(r'^\d+\.\d+\.\d+\s+[\w\u4e00-\u9fa5]+', stripped_text):
                        is_heading = True
                        heading_level = 3
                
                if is_heading:
                    # 如果当前有未完成的章节，保存它
                    if current_chapter:
                        chapters.append(current_chapter)
                    
                    # 清理标题，移除末尾的页码
                    clean_title = re.sub(r'\s*\d+$', '', stripped_text).strip()
                    
                    # 创建新章节
                    current_chapter = {
                        'title': clean_title,
                        'level': heading_level,
                        'content': '',
                        'start_paragraph': current_paragraph_index,
                        'original_text': stripped_text
                    }
                    
                    if current_paragraph_index < 50:
                        print(f"  >>> 识别为标题: {clean_title} (级别: {heading_level})")
                else:
                    # 如果是正文内容，添加到当前章节
                    if current_chapter:
                        if current_chapter['content']:
                            current_chapter['content'] += '\n'
                        current_chapter['content'] += text
                        
                        # 检查段落中是否包含图片（改进的检测逻辑）
                        has_image = False
                        
                        # 检查段落本身是否包含drawing元素
                        if paragraph._p.xml.find('w:drawing') >= 0:
                            has_image = True
                        else:
                            # 检查每个run是否包含drawing元素
                            for run in paragraph.runs:
                                if run._r.xml.find('w:drawing') >= 0:
                                    has_image = True
                                    break
                        
                        if has_image:
                            # 段落包含图片，使用唯一标记
                            image_id = f"IMAGE_MARKER_{uuid.uuid4().hex[:8]}"
                            current_chapter['content'] += f"\n[{image_id}]\n"
                            # 记录图片标记的位置
                            if 'image_markers' not in current_chapter:
                                current_chapter['image_markers'] = []
                            current_chapter['image_markers'].append(image_id)
                        
                        if current_paragraph_index < 50:
                            print(f"  >>> 添加到当前章节: {stripped_text[:50]}...")
                
                # 检查当前段落之后是否有表格需要处理
                # 更精确的表格定位：检查表格是否在当前段落附近
                # 在docx中，表格是与段落同级的元素，所以我们需要根据段落位置推断表格位置
                # 简单的策略：假设表格在其实际位置的段落之后被处理
                if current_table_index < total_tables:
                    # 这里使用一个简单的策略：每15个段落检查一次表格
                    # 在实际应用中，可能需要更复杂的逻辑来确定表格的精确位置
                    if current_paragraph_index > 0 and current_paragraph_index % 15 == 0:
                        if current_chapter:
                            table = tables[current_table_index]
                            table_md = self._convert_table_to_markdown(table)
                            
                            if table_md:
                                if current_chapter['content']:
                                    current_chapter['content'] += '\n\n'
                                current_chapter['content'] += table_md
                                current_chapter['content'] += '\n\n'
                                
                            current_table_index += 1
                
                current_paragraph_index += 1
            
            # 处理剩余的表格
            while current_table_index < total_tables:
                if current_chapter:
                    table = tables[current_table_index]
                    table_md = self._convert_table_to_markdown(table)
                    
                    if table_md:
                        if current_chapter['content']:
                            current_chapter['content'] += '\n\n'
                        current_chapter['content'] += table_md
                        current_chapter['content'] += '\n\n'
                current_table_index += 1
            
            # 保存最后一个章节
            if current_chapter:
                current_chapter['end_paragraph'] = total_paragraphs - 1
                chapters.append(current_chapter)
                
        except Exception as e:
            print(f"提取正文内容失败: {str(e)}")
            import traceback
            traceback.print_exc()
        
        return chapters
    
    def _convert_table_to_markdown(self, table) -> str:
        """将表格转换为Markdown格式
        
        Args:
            table: 表格对象
            
        Returns:
            Markdown格式的表格
        """
        if not table or not hasattr(table, 'rows'):
            return ''
            
        markdown_table = []
        
        # 转换表头
        if table.rows:
            header_row = table.rows[0]
            header_cells = [cell.text.strip() for cell in header_row.cells]
            markdown_table.append('| ' + ' | '.join(header_cells) + ' |')
            
            # 添加分隔线
            separator = ['---' for _ in header_cells]
            markdown_table.append('| ' + ' | '.join(separator) + ' |')
            
            # 转换表格内容
            for row in table.rows[1:]:
                cells = [cell.text.strip() for cell in row.cells]
                markdown_table.append('| ' + ' | '.join(cells) + ' |')
        
        return '\n'.join(markdown_table) if markdown_table else ''
        
    def extract_all_text(self) -> str:
        """提取文档所有文本内容
        
        Returns:
            文档所有文本
        """
        if not self.document:
            return ''
        
        try:
            text = ''
            for paragraph in self.document.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text.strip() + '\n'
            return text
        except Exception as e:
            print(f"提取文本失败: {str(e)}")
            return ''
    
    def extract_images(self, output_dir: str) -> List[Dict[str, Any]]:
        """提取文档中的图片并保存到指定目录
        
        Args:
            output_dir: 图片输出目录
            
        Returns:
            图片信息列表，包含图片路径、原始文件名等信息
        """
        if not self.document:
            return []
        
        images = []
        
        try:
            # 确保输出目录存在
            if not os.path.exists(output_dir):
                os.makedirs(output_dir)
            
            # 提取所有图片
            image_id = 0
            
            # 遍历文档中的所有关系，查找图片
            for rel in self.document.part.rels.values():
                if "image" in rel.target_ref:
                    image_id += 1
                    
                    # 获取图片原始数据
                    image_data = rel.target_part.blob
                    
                    # 获取图片原始文件名
                    original_filename = rel.target_part.filename
                    
                    # 生成新的文件名
                    file_ext = os.path.splitext(original_filename)[1]
                    if not file_ext:
                        # 如果没有扩展名，默认使用.png
                        file_ext = '.png'
                    file_name = f"image_{image_id}_{uuid.uuid4().hex[:8]}{file_ext}"
                    file_path = os.path.join(output_dir, file_name)
                    
                    # 保存图片
                    with open(file_path, 'wb') as f:
                        f.write(image_data)
                        
                    # 获取图片尺寸（需要PIL支持）
                    width = height = None
                    try:
                        from PIL import Image
                        with Image.open(file_path) as img:
                            width, height = img.size
                    except ImportError:
                        pass
                    except Exception as e:
                        print(f"获取图片尺寸失败: {str(e)}")
                    
                    images.append({
                        'path': file_path,
                        'filename': file_name,
                        'width': width,
                        'height': height,
                        'original_name': original_filename,
                        'id': image_id  # 添加图片ID
                    })
                    
        except Exception as e:
            print(f"提取图片失败: {str(e)}")
        
        return images


def process_docx(file_path: str, image_output_dir: str = None) -> Dict[str, Any]:
    """处理.docx文件的主要函数
    
    Args:
        file_path: .docx文件路径
        image_output_dir: 图片输出目录（可选）
        
    Returns:
        处理结果，包含目录、正文内容、图片信息和匹配结果
    """
    processor = DocxProcessor()
    result = {
        'toc': [],
        'chapters': [],
        'all_text': '',
        'images': [],
        'matched_chapters': []  # 存储目录与章节匹配的结果
    }
    
    try:
        # 打开文档
        if processor.open_document(file_path):
            # 提取目录
            result['toc'] = processor.extract_toc()
            
            # 提取章节内容
            raw_chapters = processor.extract_content_by_headings()
            
            # 过滤掉重复的章节
            filtered_chapters = []
            
            # 记录已经处理过的章节标题
            processed_titles = set()
            
            for chapter in raw_chapters:
                # 清理标题用于去重
                cleaned_title = re.sub(r'\s{2,}', ' ', chapter['title'])
                cleaned_title = re.sub(r'\s*\d+$', '', cleaned_title).strip()
                
                # 跳过已经处理过的标题
                if cleaned_title in processed_titles:
                    continue
                
                # 保留所有章节，不进行内容长度过滤
                filtered_chapters.append(chapter)
                processed_titles.add(cleaned_title)
            
            # 只保留第4章作为唯一章节，其余内容合并到这个章节中
            final_chapters = []
            
            # 先找出第4章的大标题
            chapter4 = None
            for chapter in filtered_chapters:
                title = chapter['title']
                level = chapter['level']
                
                # 找到第4章的大标题
                if level == 1 and '第4章' in title:
                    chapter4 = chapter
                    break
            
            if chapter4:
                # 将其他所有内容合并到第4章中
                all_content = chapter4['content']
                
                # 遍历所有章节，将非第4章大标题的内容合并
                for chapter in filtered_chapters:
                    if chapter != chapter4:
                        # 如果是子标题，添加Markdown格式的标题
                        if chapter['level'] == 2:
                            all_content += f"\n## {chapter['title']}\n"
                        elif chapter['level'] == 3:
                            all_content += f"\n### {chapter['title']}\n"
                        # 添加章节内容
                        if chapter['content']:
                            all_content += f"{chapter['content']}\n"
                
                # 更新第4章的内容
                chapter4['content'] = all_content
                final_chapters = [chapter4]
            else:
                # 如果没有找到第4章，保留所有章节
                final_chapters = filtered_chapters
                
            # 按开始段落排序
            final_chapters.sort(key=lambda x: x['start_paragraph'])
            
            result['chapters'] = final_chapters
            
            # 提取所有文本
            result['all_text'] = processor.extract_all_text()
            
            # 处理章节内容中的代码块和练习题
            for chapter in final_chapters:
                content = chapter['content']
                
                # 识别并处理代码块（以"# 例X.X："开头的内容）
                # 正则表达式匹配：# 例4.1：语句执行顺序 后面跟随的代码
                
                # 代码块模式：匹配以"# 例X.X："开头的行，直到遇到下一个标题或结束
                code_block_pattern = re.compile(r'(#\s*例\s*\d+\.\d+\s*：.*?)(?=(#\s*例\s*\d+\.\d+\s*：|##\s|###\s|思考与练习|$))', re.DOTALL)
                
                # 替换代码块为Jupyter代码框格式
                def replace_code_block(match):
                    code_content = match.group(1).strip()
                    return f"\n```python\n{code_content}\n```\n"
                
                content = code_block_pattern.sub(replace_code_block, content)
                
                # 识别并处理练习题（思考与练习部分）
                practice_pattern = re.compile(r'(思考与练习.*?)(?=($|#\s*例\s*\d+\.\d+\s*：))', re.DOTALL)
                
                def replace_practice(match):
                    practice_content = match.group(1).strip()
                    return f"\n## 练习题\n```python\n{practice_content}\n```\n"
                
                content = practice_pattern.sub(replace_practice, content)
                
                # 更新章节内容
                chapter['content'] = content
            
            # 提取图片
            if image_output_dir:
                result['images'] = processor.extract_images(image_output_dir)
                
                # 不直接嵌入图片引用，保留IMAGE_MARKER标记，让后续处理能够正确替换
            
            # 匹配目录与章节
            result['matched_chapters'] = _match_toc_with_chapters(result['toc'], result['chapters'])
    finally:
        # 确保关闭文档
        processor.close_document()
    
    return result


def _match_toc_with_chapters(toc_items: List[Dict[str, Any]], chapters: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    """将目录项与章节内容进行精确匹配
    
    Args:
        toc_items: 目录项列表
        chapters: 章节内容列表
        
    Returns:
        匹配后的章节列表，包含目录信息
    """
    matched_chapters = []
    
    # 创建章节标题映射表
    chapter_title_map = {}
    for chapter in chapters:
        # 标准化标题（去除标点、空格等）
        normalized_title = _normalize_title(chapter['title'])
        chapter_title_map[normalized_title] = chapter
        # 也添加原始标题作为键
        chapter_title_map[chapter['title']] = chapter
    
    # 为每个目录项找到对应的章节
    for toc_item in toc_items:
        normalized_toc_title = _normalize_title(toc_item['title'])
        
        # 尝试直接匹配标题
        matched_chapter = None
        if normalized_toc_title in chapter_title_map:
            matched_chapter = chapter_title_map[normalized_toc_title]
        elif toc_item['title'] in chapter_title_map:
            matched_chapter = chapter_title_map[toc_item['title']]
        else:
            # 尝试模糊匹配
            for normalized_chapter_title, chapter in chapter_title_map.items():
                if _is_similar_title(normalized_toc_title, normalized_chapter_title):
                    matched_chapter = chapter
                    break
        
        if matched_chapter:
            # 合并目录信息到章节中
            matched_chapter_with_toc = matched_chapter.copy()
            matched_chapter_with_toc['toc_info'] = toc_item
            matched_chapter_with_toc['toc_title'] = toc_item['title']
            matched_chapter_with_toc['toc_level'] = toc_item['level']
            matched_chapter_with_toc['toc_page_num'] = toc_item['page_num']
            matched_chapters.append(matched_chapter_with_toc)
        else:
            # 如果没有匹配到，创建一个空章节
            matched_chapters.append({
                'title': toc_item['title'],
                'level': toc_item['level'],
                'content': '',
                'toc_info': toc_item,
                'is_missing': True
            })
    
    # 如果有章节没有在目录中找到，也添加到结果中
    for chapter in chapters:
        normalized_chapter_title = _normalize_title(chapter['title'])
        found = False
        for matched in matched_chapters:
            if 'toc_info' not in matched and 'title' in matched:
                if _normalize_title(matched['title']) == normalized_chapter_title:
                    found = True
                    break
        
        if not found:
            matched_chapters.append(chapter)
    
    # 按级别排序
    matched_chapters.sort(key=lambda x: (x.get('level', 999), x.get('start_paragraph', 999)))
    
    return matched_chapters


def _normalize_title(title: str) -> str:
    """标准化标题，用于匹配
    
    Args:
        title: 原始标题
        
    Returns:
        标准化后的标题
    """
    if not title:
        return ""
    
    # 转换为小写
    title = title.lower()
    
    # 去除标点符号和空白字符
    # 使用更兼容的正则表达式，避免使用Unicode属性转义
    title = re.sub(r'[\.,;:"!@#$%^&*()\-_=+\[\]{}\\|<>?/\\`~\s]+', ' ', title)
    
    # 去除多余空格
    title = re.sub(r'\s+', ' ', title)
    
    # 去除首尾空格
    return title.strip()


def _is_similar_title(title1: str, title2: str) -> bool:
    """判断两个标题是否相似
    
    Args:
        title1: 第一个标题
        title2: 第二个标题
        
    Returns:
        如果相似返回True，否则返回False
    """
    if not title1 or not title2:
        return False
    
    # 使用简单的字符串相似度判断
    import difflib
    similarity = difflib.SequenceMatcher(None, title1, title2).ratio()
    
    # 如果相似度超过80%，认为是相似的
    return similarity > 0.8


if __name__ == "__main__":
    # 测试代码
    import sys
    if len(sys.argv) < 2:
        print("用法: python docx_processor.py <docx文件路径>")
        sys.exit(1)
    
    docx_path = sys.argv[1]
    if not os.path.exists(docx_path):
        print(f"文件不存在: {docx_path}")
        sys.exit(1)
    
    print(f"开始处理文档: {docx_path}")
    
    # 处理文档
    with tempfile.TemporaryDirectory() as temp_dir:
        result = process_docx(docx_path, image_output_dir=temp_dir)
    
    print(f"目录项数: {len(result['toc'])}")
    print(f"章节数: {len(result['chapters'])}")
    print(f"匹配后的章节数: {len(result['matched_chapters'])}")
    print(f"提取的图片数: {len(result['images'])}")
    
    # 打印目录
    print("\n目录:")
    for i, item in enumerate(result['toc'][:5]):
        print(f"{i+1}. [{item['level']}] {item['title']} (第{item['page_num']}页)")
    
    # 打印章节
    print("\n章节:")
    for i, chapter in enumerate(result['chapters'][:3]):
        print(f"\n章节 {i+1}:")
        print(f"标题: {chapter['title']}")
        print(f"级别: {chapter['level']}")
        content_preview = chapter['content'][:100] + '...' if len(chapter['content']) > 100 else chapter['content']
        print(f"内容预览: {content_preview}")