#!/usr/bin/env python3
# -*- coding: utf-8 -*-

import os
import sys

# 添加项目根目录和apps目录到Python路径
sys.path.append(os.path.dirname(os.path.abspath(__file__)))
sys.path.append(os.path.join(os.path.dirname(os.path.abspath(__file__)), 'apps'))

from docx import Document

def test_docx_structure():
    """测试DOCX文档结构，了解如何访问表格和内联图片"""
    # 测试文件路径
    test_file_path = r"d:\liufyData\CodeBook_max\《人工智能基础》讲义——中.docx"
    
    if not os.path.exists(test_file_path):
        print(f"测试文件不存在: {test_file_path}")
        return
    
    print(f"测试文件: {test_file_path}")
    
    # 打开文档
    doc = Document(test_file_path)
    
    # 检查文档结构
    print(f"\n=== 文档结构分析 ===")
    print(f"段落数量: {len(doc.paragraphs)}")
    print(f"表格数量: {len(doc.tables)}")
    
    # 查看前5个表格的结构
    for i, table in enumerate(doc.tables[:5]):
        print(f"\n表格 {i+1}:")
        print(f"  行数: {len(table.rows)}")
        print(f"  列数: {len(table.columns if hasattr(table, 'columns') else table.rows[0].cells)}")
        
        # 打印表格前2行内容
        for j, row in enumerate(table.rows[:2]):
            cells = row.cells
            cell_texts = [cell.text.strip() for cell in cells]
            print(f"  行 {j+1}: {cell_texts}")
    
    # 查看段落中的内联图片
    print(f"\n=== 内联图片分析 ===")
    for i, paragraph in enumerate(doc.paragraphs[:20]):  # 只查看前20个段落
        runs = paragraph.runs
        print(f"\n段落 {i+1} 包含 {len(runs)} 个runs")
        
        for j, run in enumerate(runs):
            print(f"  Run {j+1}: 文本='{run.text}'")
            
            # 检查run中是否包含图片
            if run.element.rPr is not None:
                for child in run.element.iterchildren():
                    if 'drawing' in child.tag:
                        print(f"  → Run {j+1} 包含图片！")

if __name__ == "__main__":
    test_docx_structure()